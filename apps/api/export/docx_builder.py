from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO

from .prosemirror_to_docx import render_prosemirror

MEMBERSHIP_TIERS = [
    {
        "name": "Gold Membership", "price": "$109/month",
        "benefits": [
            "30% off Pure Encapsulations supplements",
            "20% off Thorne vitamins",
            "30% off IV infusions plus rotating 50% off deals",
            "Access to premium peptides including GH, NAD+, Mounjaro, and Glutathione",
            "One free lab panel per year",
            "Priority scheduling and concierge text access during business hours",
            "Discounts on prescription medications with first priority",
        ],
    },
    {
        "name": "Hormone Optimization Membership", "price": "$75/month",
        "benefits": [
            "Personalized hormone treatment plan including BHRT, peptides, and labs",
            "30% off Pure Encapsulations hormone, adrenal, and metabolic support",
            "20% off Thorne vitamins",
            "20% off IV infusions plus rotating 50% off deals",
            "Discounts on peptides including Mounjaro, Ozempic, NAD+, GH, and Glutathione",
            "Discounts on prescription medications including hormone therapies",
            "10 yearly visits or texts with prescriptions",
        ],
    },
    {
        "name": "Silver Membership", "price": "$49.99/month",
        "benefits": [
            "20% off Pure Encapsulations supplements",
            "10% off Thorne vitamins",
            "10% off IV infusions",
            "Discounts on prescription medications",
            "Option to upgrade anytime to higher tiers",
            "Annual lab panel",
            "5 yearly visits or texts with prescriptions",
        ],
    },
]

HEADER_FIELDS = [
    ("Patient Name", "patient_name"),
    ("DOB", "dob"),
    ("Visit Date", "visit_date"),
    ("Primary Focus", "primary_focus"),
    ("Current Regimen", "current_regimen"),
    ("Compared To", "compared_to"),
    ("Protocol Prepared By", "protocol_prepared_by"),
    ("Source Documents", "source_documents"),
]


def _add_header_table(doc, header_data: dict):
    table = doc.add_table(rows=len(HEADER_FIELDS), cols=2)
    table.style = "Table Grid"
    for i, (label, key) in enumerate(HEADER_FIELDS):
        cells = table.rows[i].cells
        run = cells[0].paragraphs[0].add_run(label)
        run.bold = True
        cells[1].text = header_data.get(key, "") or ""
    doc.add_paragraph()


def _add_version_label(doc, label: str):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x3C, 0x34, 0x89)
    doc.add_paragraph()


def _add_section(doc, display_name: str, content: dict | None):
    heading = doc.add_paragraph()
    run = heading.add_run(display_name)
    run.bold = True
    run.font.size = Pt(12)
    render_prosemirror(doc, content)
    doc.add_paragraph()


def _add_supplement_table(doc, rows: list[dict]):
    heading = doc.add_paragraph()
    run = heading.add_run("FINAL SUPPLEMENTATION LIST")
    run.bold = True
    run.font.size = Pt(13)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, label in enumerate(["#", "Supplement", "Dose", "Instructions / Purpose"]):
        run = table.rows[0].cells[i].paragraphs[0].add_run(label)
        run.bold = True

    for i, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        cells[0].text = str(i)
        cells[1].text = row.get("name", "")
        cells[2].text = row.get("dose", "")
        instr = " - ".join(filter(None, [row.get("timing", ""), row.get("purpose", "")]))
        cells[3].text = instr
    doc.add_paragraph()


def _add_membership_table(doc):
    heading = doc.add_paragraph()
    run = heading.add_run("IQMD MEMBERSHIP LEVELS")
    run.bold = True
    run.font.size = Pt(13)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, label in enumerate(["Membership", "Price", "Included Benefits"]):
        run = table.rows[0].cells[i].paragraphs[0].add_run(label)
        run.bold = True

    for tier in MEMBERSHIP_TIERS:
        cells = table.add_row().cells
        cells[0].text = tier["name"]
        cells[1].text = tier["price"]
        cells[2].text = "\n".join(f"- {b}" for b in tier["benefits"])


def _add_draft_watermark(doc):
    header = doc.sections[0].header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DRAFT")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0xE5, 0xB8, 0xB8)


def _add_page_numbers(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()

    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = "PAGE"
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)


def build_protocol_docx(
    document, physician_sections, patient_sections, supplement_rows, header_data
) -> bytes:
    doc = DocxDocument()

    if document.status.value == "draft":
        _add_draft_watermark(doc)
    _add_page_numbers(doc)

    doc.add_heading("IQMD FUNCTIONAL MEDICINE PROTOCOL", level=1)
    sub = doc.add_paragraph()
    sub.add_run("Prepared using IQMD Master Template 2026").italic = True
    doc.add_paragraph()

    _add_header_table(doc, header_data)

    _add_version_label(doc, "PHYSICIAN VERSION")
    for s in physician_sections:
        if s.field_type.value == "auto_fill":
            continue
        _add_section(doc, s.display_name, document.sections_data.get(s.section_key))

    doc.add_page_break()
    _add_header_table(doc, header_data)  # repeated per locked-use rule

    _add_version_label(doc, "PATIENT VERSION")
    for s in patient_sections:
        if s.field_type.value == "auto_fill":
            continue
        _add_section(doc, s.display_name, document.sections_data.get(s.section_key))

    doc.add_page_break()
    _add_supplement_table(doc, supplement_rows)
    _add_membership_table(doc)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
