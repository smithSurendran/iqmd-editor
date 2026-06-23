from docx.shared import Pt


def _apply_marks(run, marks: list[dict] | None):
    for mark in marks or []:
        t = mark.get("type")
        if t == "bold":
            run.bold = True
        elif t == "italic":
            run.italic = True
        elif t == "underline":
            run.underline = True


def _add_text_runs(paragraph, content: list[dict]):
    for node in content or []:
        if node.get("type") == "text":
            run = paragraph.add_run(node.get("text", ""))
            _apply_marks(run, node.get("marks"))


def _add_list(doc, list_node: dict, style: str):
    for item in list_node.get("content", []):
        for para in item.get("content", []):
            if para.get("type") == "paragraph":
                p = doc.add_paragraph(style=style)
                _add_text_runs(p, para.get("content", []))


def render_prosemirror(doc, pm_json: dict | None):
    """Renders TipTap/ProseMirror JSON into the docx document in place."""
    if not pm_json or not pm_json.get("content"):
        p = doc.add_paragraph()
        run = p.add_run("-")
        run.italic = True
        return

    for node in pm_json["content"]:
        node_type = node.get("type")
        if node_type == "paragraph":
            if not node.get("content"):
                doc.add_paragraph()
                continue
            p = doc.add_paragraph()
            _add_text_runs(p, node.get("content", []))
        elif node_type == "heading":
            p = doc.add_paragraph()
            _add_text_runs(p, node.get("content", []))
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)
        elif node_type == "bulletList":
            _add_list(doc, node, "List Bullet")
        elif node_type == "orderedList":
            _add_list(doc, node, "List Number")
